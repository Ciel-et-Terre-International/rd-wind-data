import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def set_paragraph_font(paragraph, size=10):
    """Applique une taille de police uniforme à tous les runs d’un paragraphe."""
    for run in paragraph.runs:
        run.font.size = Pt(size)


def insert_section_title(doc, title, level=1):
    """
    Insère un titre de section avec le style de titre (Heading 1/2...)
    du template Word. Exemple : "1 | Stations context".
    """
    heading = doc.add_heading(title, level=level)
    set_paragraph_font(heading, size=12)


def insert_paragraph(doc, text):
    """Insère un paragraphe aligné à gauche avec une petite taille de police."""
    para = doc.add_paragraph(text)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_font(para, size=9)


def insert_image_if_exists(doc, path, width=6.0):
    """
    Insère une image si elle existe, avec un saut de ligne après.
    width est en pouces (Inches).
    """
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        # petit espace après l'image
        spacer = doc.add_paragraph()
        set_paragraph_font(spacer, size=6)


def insert_table_from_csv(doc, csv_path, title=None):
    """
    Insère une table à partir d’un CSV (en conservant l’ordre des colonnes).
    Si le fichier n’existe pas ou est vide, ne fait rien.
    """
    if not os.path.exists(csv_path):
        return

    df = pd.read_csv(csv_path)
    if df.empty:
        return

    if title:
        insert_section_title(doc, title, level=2)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    # En-têtes
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Lignes
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = "" if pd.isna(val) else str(val)

    # petit espace après la table
    spacer = doc.add_paragraph()
    set_paragraph_font(spacer, size=6)


def generate_report(site_data, output_folder="data"):
    """
    Génère le rapport Word pour un site, en utilisant le template Word
    (template.docx) pour la mise en page (logo, bandeau, table des matières).
    Le contenu textuel et les tableaux/images suivent la structure du rapport
    ITAIPU fourni en exemple.

    Hypothèse :
      - Le template contient déjà la page de garde, en-têtes/pieds de page,
        et éventuellement la table des matières. On commence donc directement
        par les sections 1, 2, 3, etc.
    """
    # Nom de site au format "REF_SITE" (ex : WBR043_ITAIPU)
    site_name = f"{site_data['reference']}_{site_data['name']}"

    site_folder = os.path.join(output_folder, site_name)
    figures_dir = os.path.join(site_folder, "figures_and_tables")

    output_docx = os.path.join(site_folder, "report", f"fiche_{site_name}.docx")
    os.makedirs(os.path.dirname(output_docx), exist_ok=True)

    # Chargement du template Word
    template_path = os.path.join(os.getcwd(), "template.docx")
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        print("template.docx introuvable. Utilisation d’un document Word vierge.")
        doc = Document()

    # =========================================================================
    # 1 | Stations context
    # =========================================================================
    insert_section_title(doc, "1 | Stations context", level=1)
    insert_paragraph(
        doc,
        (
            "This table provides information about the meteorological stations used "
            "(IDs, names, distances, coordinates)."
        ),
    )

    # On essaie d’abord dans figures_and_tables, puis à la racine du site.
    stations_csv_candidates = [
        os.path.join(figures_dir, "stations_context.csv"),
        os.path.join(site_folder, "stations_context.csv"),
    ]
    for path in stations_csv_candidates:
        if os.path.exists(path):
            insert_table_from_csv(doc, path)
            break

    # =========================================================================
    # 2 | Datas quality
    # =========================================================================
    insert_section_title(doc, "2 | Datas quality", level=1)
    insert_paragraph(
        doc,
        (
            "This table shows the number of days of available data, the study period, "
            "and the coverage rate for average wind and gusts for each source used."
        ),
    )
    insert_table_from_csv(doc, os.path.join(figures_dir, "resume_qualite.csv"))

    # =========================================================================
    # 3 | Statistics description
    # =========================================================================
    insert_section_title(doc, "3 | Statistics description", level=1)
    insert_paragraph(
        doc,
        (
            "Distribution of average wind speeds by source. Average, standard deviation, "
            "percentiles, min/max. All speeds are expressed in m/s."
        ),
    )
    insert_table_from_csv(doc, os.path.join(figures_dir, "stats_windspeed_mean.csv"))

    # =========================================================================
    # 4 | Histograms
    # =========================================================================
    insert_section_title(doc, "4 | Histograms", level=1)
    insert_paragraph(
        doc,
        (
            "Histograms showing the distribution of average and gust wind speeds by "
            "source, with visualisation of the regulatory threshold (Building Code)."
        ),
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "histogrammes_windspeed_mean.png"),
        width=6.0,
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "histogrammes_windspeed_gust.png"),
        width=6.0,
    )

    # =========================================================================
    # 5 | Extreme values analysis
    # =========================================================================
    insert_section_title(doc, "5 | Extreme values analysis", level=1)
    insert_paragraph(
        doc,
        (
            "Visualisation of outliers via box plots and histograms. These values may "
            "influence the analysis of gusts or extremes. The summary tables also "
            "indicate the number of days above the Building Code thresholds."
        ),
    )

    # Boxplots + histogrammes de jours extrêmes
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "boxplot_windspeed_mean.png"),
        width=6.0,
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "outliers_hist_windspeed_mean.png"),
        width=5.0,
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "boxplot_windspeed_gust.png"),
        width=6.0,
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "outliers_hist_windspeed_gust.png"),
        width=5.0,
    )

    # Tableaux jours extrêmes – vent moyen / rafales (résumé + par année)
    insert_paragraph(doc, "Summary of days above Building Code thresholds (mean wind).")
    insert_table_from_csv(
        doc, os.path.join(figures_dir, "vent_moyen_extremes_resume.csv")
    )
    insert_table_from_csv(
        doc, os.path.join(figures_dir, "vent_moyen_extremes_par_an.csv")
    )

    insert_paragraph(doc, "Summary of days above Building Code thresholds (gust wind).")
    insert_table_from_csv(
        doc, os.path.join(figures_dir, "rafales_extremes_resume.csv")
    )
    insert_table_from_csv(
        doc, os.path.join(figures_dir, "rafales_extremes_par_an.csv")
    )

    # =========================================================================
    # 6 | Directional analysis – Compass roses
    # =========================================================================
    insert_section_title(doc, "6 | Directional analysis – Compass roses", level=1)
    insert_paragraph(
        doc,
        (
            "Visualisation of prevailing wind directions by source, including "
            "maximum daily wind speeds and occurrence frequency for each 20° sector."
        ),
    )

    if os.path.isdir(figures_dir):
        for f in sorted(os.listdir(figures_dir)):
            if f.startswith("rose_max_windspeed_") and f.endswith(".png"):
                insert_image_if_exists(doc, os.path.join(figures_dir, f), width=5.5)
        for f in sorted(os.listdir(figures_dir)):
            if f.startswith("rose_frequency_") and f.endswith(".png"):
                insert_image_if_exists(doc, os.path.join(figures_dir, f), width=5.5)

    # =========================================================================
    # 7 | Time series
    # =========================================================================
    insert_section_title(doc, "7 | Time series", level=1)
    insert_paragraph(
        doc,
        (
            "View complete time series for average wind and gusts. "
            "This helps to identify major trends and potential extreme peaks."
        ),
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "time_series_windspeed_mean.png"),
        width=6.5,
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "time_series_windspeed_gust.png"),
        width=6.5,
    )

    # =========================================================================
    # 8 | Analysis of levels for a 50-year return period
    # =========================================================================
    insert_section_title(doc, "8 | Analysis of levels for a 50-year return period", level=1)
    insert_paragraph(
        doc,
        (
            "Estimation of average wind speeds and gusts with a 50-year return period, "
            "according to Gumbel's law. Comparison with regulatory thresholds."
        ),
    )
    insert_table_from_csv(doc, os.path.join(figures_dir, "return_period_50y.csv"))

    # =========================================================================
    # Sauvegarde
    # =========================================================================
    doc.save(output_docx)
    print(f"Rapport DOCX généré : {output_docx}")
