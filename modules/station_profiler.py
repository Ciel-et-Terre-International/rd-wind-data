import os
import requests
import pandas as pd
from docx import Document
from docx.shared import Pt

def get_elevation(lat, lon):
    try:
        url = f"https://api.open-elevation.com/api/v1/lookup?locations={lat},{lon}"
        r = requests.get(url)
        if r.status_code == 200:
            return r.json()["results"][0]["elevation"]
    except Exception as e:
        print(f"Erreur Open-Elevation : {e}")
    return None

def estimate_roughness(terrain_context):
    table = {
        "mer": 0.0002,
        "plaine": 0.03,
        "campagne": 0.1,
        "forêt": 0.4,
        "urbain": 1.0,
        "centre-ville": 2.0,
        "montagne": 0.3,
        "inconnu": ""
    }
    return table.get(terrain_context.lower(), "")

def generate_station_csv(site_name, site_info, station1, station2, noaa1=None, noaa2=None):
    site_folder = os.path.join("data", site_info["reference"] + "_" + site_name)
    os.makedirs(site_folder, exist_ok=True)

    station_rows = []
    for i, station in enumerate([station1, station2], 1):
        lat = station.get("latitude")
        lon = station.get("longitude")
        altitude = get_elevation(lat, lon)
        terrain = "plaine" if altitude and altitude < 300 else "montagne" if altitude and altitude > 1000 else "campagne"

        row = {
            "site_name": site_name,
            "source": f"meteostat{i}",
            "station_id": station.get("id"),
            "station_name": station.get("name"),
            "latitude": lat,
            "longitude": lon,
            "distance_km": station.get("distance_km"),
            "altitude_m": altitude,
            "anemometer_height_m": 10,
            "station_type": "unknown",
            "start_date": None,
            "end_date": None,
            "terrain_context": terrain,
            "roughness_estimate": estimate_roughness(terrain),
            "data_coverage_percent": None
        }

        csv_path = os.path.join(site_folder, f"meteostat{i}_{site_name}.csv")
        if os.path.exists(csv_path):
            df = pd.read_csv(csv_path)
            total_days = pd.date_range(site_info["start"], site_info["end"]).shape[0]
            valid_days = df.dropna(subset=["windspeed_mean", "wind_direction"]).shape[0]
            row["data_coverage_percent"] = round(100 * valid_days / total_days, 1)

        station_rows.append(row)

    for i, station in enumerate([noaa1, noaa2], 1):
        if not station:
            continue
        lat = station.get("lat")
        lon = station.get("lon")
        altitude = station.get("elev") or get_elevation(lat, lon)
        terrain = "plaine" if altitude and altitude < 300 else "montagne" if altitude and altitude > 1000 else "campagne"

        row = {
            "site_name": site_name,
            "source": f"noaa_station{i}",
            "station_id": f"{station['usaf']}-{station['wban']}",
            "station_name": station.get("name"),
            "latitude": lat,
            "longitude": lon,
            "distance_km": round(station.get("distance_km", 0), 2),
            "altitude_m": altitude,
            "anemometer_height_m": station.get("anemometer_height", 10),
            "station_type": station.get("station_type", "unknown"),
            "start_date": station.get("begin"),
            "end_date": station.get("end"),
            "terrain_context": terrain,
            "roughness_estimate": estimate_roughness(terrain),
            "data_coverage_percent": None
        }

        csv_path = os.path.join(site_folder, f"noaa_station{i}_{site_name}.csv")
        if os.path.exists(csv_path):
            df = pd.read_csv(csv_path)
            total_days = pd.date_range(site_info["start"], site_info["end"]).shape[0]
            valid_days = df.dropna(subset=["windspeed_mean", "wind_direction"]).shape[0]
            row["data_coverage_percent"] = round(100 * valid_days / total_days, 1)

        station_rows.append(row)

    output_csv = os.path.join(site_folder, f"stations_{site_name}.csv")
    pd.DataFrame(station_rows).to_csv(output_csv, index=False)
    print(f"Fichier CSV des stations généré : {output_csv}")
    return station_rows

def generate_station_docx(site_name, station_data_list, output_path):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()
    doc.add_heading(f"Fiche Stations – {site_name}", level=1)

    for data in station_data_list:
        doc.add_heading(f"Station : {data['station_name']} ({data['station_id']})", level=2)
        doc.add_paragraph(f"Source : {data['source']}")
        doc.add_paragraph(f"Coordonnées : {data['latitude']} / {data['longitude']}")
        doc.add_paragraph(f"Distance au site : {data['distance_km']} km")
        doc.add_paragraph(f"Altitude : {data['altitude_m']} m")
        doc.add_paragraph(f"Hauteur anémomètre : {data['anemometer_height_m']} m")
        doc.add_paragraph(f"Type de station : {data['station_type']}")
        doc.add_paragraph(f"Période mesurée : {data['start_date']} → {data['end_date']}")
        doc.add_paragraph(f"Complétion des données : {data['data_coverage_percent']} %")
        doc.add_paragraph(f"Contexte terrain : {data['terrain_context']}")
        doc.add_paragraph(f"Rugosité estimée : {data['roughness_estimate']}")

    doc.save(output_path)
    print(f"Fichier DOCX des stations généré : {output_path}")
    return output_path
