import os
import pandas as pd
from pathlib import Path
from collections import defaultdict
from docx import Document
from docx.shared import Pt

# === CONFIGURATION ===
DATA_DIR = Path("data")
SITE_CONFIG_FILE = Path("modele_sites.csv")
OUTPUT_FILE = Path("analyse_globale_par_pays.csv")

# === CHARGEMENT DU FICHIER DE RÉFÉRENCE ===
def load_site_config():
    print("📄 Chargement du fichier modele_sites.csv...")
    df = pd.read_csv(SITE_CONFIG_FILE)
    print(f"🔢 {len(df)} sites trouvés dans le fichier de configuration.")
    return df

# === CHARGEMENT DES DONNÉES DE PÉRIODES DE RETOUR POUR UN SITE ===
def load_return_data(site_path):
    file_path = site_path / "figures_and_tables" / "return_period_50y.csv"
    if not file_path.exists():
        print(f"❌ Fichier manquant : {file_path}")
        return None
    df = pd.read_csv(file_path)

    if not {"Source", "Variable", "Return_Period_50y (m/s)"}.issubset(df.columns):
        print(f"❌ Colonnes attendues absentes dans {file_path}")
        return None

    df_mean = df[(df["Variable"] == "windspeed_mean") & (~df["Source"].str.lower().str.startswith("noaa"))].copy()
    df_mean.dropna(subset=["Return_Period_50y (m/s)", "Source"], inplace=True)
    df_mean.rename(columns={"Return_Period_50y (m/s)": "windspeed_mean_50y"}, inplace=True)

    if df_mean.empty:
        print(f"⚠️ Aucune valeur windspeed_mean valide trouvée dans {file_path}")
    else:
        print(f"📊 {len(df_mean)} lignes valides (windspeed_mean) dans {file_path}.")

    return df_mean

# === CALCUL DE L'ÉCART VS BUILDING CODE POUR CHAQUE SOURCE ===
def compare_sources_to_building_code(df, building_code_value):
    ecarts = []
    nb_depassements = 0

    for _, row in df.iterrows():
        observed = row["windspeed_mean_50y"]
        ecart = observed - building_code_value
        ecarts.append(ecart)
        if observed > building_code_value:
            nb_depassements += 1

    return ecarts, nb_depassements

# === INDICE DE FIABILITÉ (score sur 5) ===
def compute_reliability_index(df):
    score = 0

    if len(df) >= 2:
        score += 1

    if {"data_start", "data_end"}.issubset(df.columns):
        coverage_years = [
            (pd.to_datetime(r["data_end"]) - pd.to_datetime(r["data_start"])).days / 365.25
            for _, r in df.iterrows()
        ]
        if all(years >= 20 for years in coverage_years):
            score += 1

        if all(
            pd.to_datetime(r["data_start"]) <= pd.to_datetime("1990-01-01") and
            pd.to_datetime(r["data_end"]) >= pd.to_datetime("2020-01-01")
            for _, r in df.iterrows()
        ):
            score += 1

    if "distance_km" in df.columns and all(r["distance_km"] <= 25 for _, r in df.iterrows()):
        score += 1

    if "duration" in df.columns and all(pd.notna(r["duration"]) for _, r in df.iterrows()):
        score += 1

    return f"{score}/5"

# === PHRASE AUTOMATIQUE ===
def generate_analysis_sentence(country, site_name, ecarts, nb_depassements, nb_sources, building_code):
    ecart_min = min(ecarts)
    ecart_max = max(ecarts)
    ratio_depassement = nb_depassements / nb_sources if nb_sources else 0

    if nb_sources == 0:
        return "Aucune source valide disponible."

    if nb_depassements == 0:
        tendance = "aucune source n'a dépassé la valeur normative, suggérant un surdimensionnement possible"
    elif ratio_depassement >= 0.5:
        tendance = "plus de la moitié des sources dépassent le seuil, ce qui pourrait signaler une sous-estimation des normes"
    else:
        tendance = "quelques sources dépassent le seuil, ce qui reste à surveiller localement"

    return (
        f"{nb_sources} sources valides ont été analysées pour le site {site_name} ({country}). "
        f"Les vitesses de vent moyen observées (50 ans) varient entre {ecart_min:.1f} et {ecart_max:.1f} m/s par rapport à la valeur du Building Code ({building_code:.1f} m/s). "
        f"{nb_depassements} dépassement(s) ont été détectés, soit {ratio_depassement:.0%} des sources. "
        f"Conclusion : {tendance}."
    )

# === RAPPORT WORD ===
def generate_word_report(df):
    doc = Document()
    doc.add_heading("Rapport global - Analyse des écarts vent / normes", level=1)

    for country in sorted(df["pays"].unique()):
        doc.add_heading(f"{country}", level=2)
        df_country = df[df["pays"] == country]

        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Site"
        hdr_cells[1].text = "Sources"
        hdr_cells[2].text = "Écart min (m/s)"
        hdr_cells[3].text = "Écart max (m/s)"
        hdr_cells[4].text = "Dépassements"

        for _, row in df_country.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row["site"]
            row_cells[1].text = str(row["nb_sources"])
            row_cells[2].text = f"{row['ecart_min']:.1f}"
            row_cells[3].text = f"{row['ecart_max']:.1f}"
            row_cells[4].text = str(row["nb_depassements"])

        doc.add_paragraph("")
        for _, row in df_country.iterrows():
            para = doc.add_paragraph(row["phrase_analyse"])
            para.runs[0].font.size = Pt(10)

    output_path = "rapport_global.docx"
    doc.save(output_path)
    print(f"📄 Rapport Word généré : {output_path}")

# === EXCEL AVEC POURCENTAGES ===
def generate_excel_summary(df):
    output_path = "resume_analyse.xlsx"

    # Calcul des écarts en pourcentage par site
    df = df.copy()
    df["val_observee_min"] = df["building_code"] + df["ecart_min"   ]
    df["val_observee_max"] = df["building_code"] + df["ecart_max"]
    df["val_observee_moy"] = df["building_code"] + df["ecart_moyen"]

    df["écart min (%)"] = -100 * (1 - df["val_observee_min"] / df["building_code"])
    df["écart max (%)"] = -100 * (1 - df["val_observee_max"] / df["building_code"])
    df["écart moyen (%)"] = -100 * (1 - df["val_observee_moy"] / df["building_code"])

    # Suppression colonnes intermédiaires
    df.drop(columns=["val_observee_min", "val_observee_max", "val_observee_moy"], inplace=True)

    with pd.ExcelWriter(output_path, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Synthèse complète", index=False)
        for country in sorted(df["pays"].unique()):
            df[df["pays"] == country].to_excel(writer, sheet_name=country[:31], index=False)

    print(f"📊 Résumé Excel généré : {output_path}")


# === RÉSUMÉ PAR PAYS ===
def generate_country_summary(df):
    summary = []

    for country in sorted(df["pays"].unique()):
        df_country = df[df["pays"] == country]
        nb_sites = len(df_country)
        nb_sources = df_country["nb_sources"].sum()
        nb_sites_depassement = (df_country["nb_depassements"] > 0).sum()
        ratio_sites_depassement = nb_sites_depassement / nb_sites if nb_sites else 0

        avg_ecart_min = df_country["ecart_min"].mean()
        avg_ecart_max = df_country["ecart_max"].mean()
        avg_ecart_moyen = df_country["ecart_moyen"].mean()

        ecarts_min_pct, ecarts_max_pct, ecarts_moy_pct = [], [], []

        for _, row in df_country.iterrows():
            bc = row["building_code"]
            if bc != 0:
                val_min = bc + row["ecart_min"]
                val_max = bc + row["ecart_max"]
                val_moy = bc + row["ecart_moyen"]
                ecarts_min_pct.append(-100 * (1 - val_min / bc))
                ecarts_max_pct.append(-100 * (1 - val_max / bc))
                ecarts_moy_pct.append(-100 * (1 - val_moy / bc))

        pct_min = sum(ecarts_min_pct) / len(ecarts_min_pct) if ecarts_min_pct else None
        pct_max = sum(ecarts_max_pct) / len(ecarts_max_pct) if ecarts_max_pct else None
        pct_moy = sum(ecarts_moy_pct) / len(ecarts_moy_pct) if ecarts_moy_pct else None

        if nb_sites_depassement == 0:
            conclusion = "Aucun site n’a dépassé la norme. Tendance au surdimensionnement."
        elif ratio_sites_depassement >= 0.5:
            conclusion = "Plus de la moitié des sites dépassent la norme. Vigilance requise."
        else:
            conclusion = "Certains sites dépassent ponctuellement les normes."

        summary.append({
            "pays": country,
            "nb_sites": nb_sites,
            "nb_sources_total": nb_sources,
            "% sites dépassant le BC": f"{ratio_sites_depassement:.0%}",
            "écart min moyen (m/s)": f"{avg_ecart_min:.1f}",
            "écart max moyen (m/s)": f"{avg_ecart_max:.1f}",
            "écart moyen (m/s)": f"{avg_ecart_moyen:.1f}",
            "écart min (%)": f"{pct_min:.0f}%" if pct_min is not None else "N/A",
            "écart max (%)": f"{pct_max:.0f}%" if pct_max is not None else "N/A",
            "écart moyen (%)": f"{pct_moy:.0f}%" if pct_moy is not None else "N/A",
            "conclusion": conclusion
        })

    return pd.DataFrame(summary)


# === AGRÉGATION PRINCIPALE ===
def main():
    site_config = load_site_config()
    results = []

    for index, row in site_config.iterrows():
        site_name = row["name"]
        country = row["country"]
        folder_name = f"{row['reference']}_{row['name']}"
        site_path = DATA_DIR / folder_name

        print(f"\n🔍 Traitement du site {site_name} ({folder_name})...")

        building_code = row.get("building_code_windspeed_mean_50y", None)
        if pd.isna(building_code):
            print("⚠️ Valeur de Building Code manquante. Site ignoré.")
            continue

        df = load_return_data(site_path)
        if df is None or df.empty:
            print("⚠️ Données absentes ou invalides. Site ignoré.")
            continue

        ecarts, nb_depassements = compare_sources_to_building_code(df, building_code)
        nb_sources = len(ecarts)
        ecart_min = min(ecarts)
        ecart_max = max(ecarts)
        fiabilite = compute_reliability_index(df)
        phrase = generate_analysis_sentence(country, site_name, ecarts, nb_depassements, nb_sources, building_code)

        print(f"✅ {nb_sources} sources valides. Fiabilité : {fiabilite}.")

        results.append({
            "pays": country,
            "site": site_name,
            "nb_sources": nb_sources,
            "ecart_min": ecart_min,
            "ecart_max": ecart_max,
            "nb_depassements": nb_depassements,
            "fiabilite": fiabilite,
            "building_code": building_code,
            "ecart_moyen": (ecart_min + ecart_max) / 2,
            "phrase_analyse": phrase
        })

    df_out = pd.DataFrame(results)
    df_out.to_csv(OUTPUT_FILE, index=False)
    df_resume_pays = generate_country_summary(df_out)
    df_resume_pays.to_excel("resume_pays.xlsx", index=False)
    print("📊 Résumé par pays généré : resume_pays.xlsx")
    print(f"\n📤 Analyse globale exportée vers {OUTPUT_FILE} (Total : {len(results)} sites analysés)")

    generate_word_report(df_out)
    generate_excel_summary(df_out)

if __name__ == "__main__":
    main()